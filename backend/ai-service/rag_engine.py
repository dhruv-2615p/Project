"""
RAG Engine with ChromaDB + Google Embeddings
Based on LangChain approach - uses Google's embedding model directly
No external model downloads needed - uses same API key as Gemini
"""

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

import os
import re
import time
from dotenv import load_dotenv
from typing import List, Dict, Tuple

# Load environment variables
load_dotenv()

# Persistent ChromaDB directory
CHROMA_PERSIST_DIR = os.path.join(os.path.dirname(__file__), ".chromadb")


class DocumentProcessor:
    """Process various document formats into text"""
    
    def load_document(self, file_path: str) -> str:
        """Load document content based on file type"""
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.txt' or ext == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif ext == '.pdf':
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text
                except ImportError:
                    print("pypdf not installed, skipping PDF")
                    return ""
            
            elif ext == '.docx':
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return text
                except ImportError:
                    print("python-docx not installed, skipping DOCX")
                    return ""
            
            else:
                print(f"Unsupported file type: {ext}")
                return ""
                
        except Exception as e:
            print(f"Error loading {file_path}: {e}")
            return ""


class RAGEngine:
    """
    RAG Engine with ChromaDB + Google Embeddings
    Uses Google's embedding model (models/embedding-001) - same API key as Gemini
    """
    
    def __init__(self):
        # Get API key
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")
        
        os.environ["GOOGLE_API_KEY"] = api_key
        
        # Initialize Gemini LLM
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            temperature=0.7,
            max_output_tokens=4096
        )
        
        # Initialize Google Embeddings (new model name: gemini-embedding-001)
        print("Initializing Google Embeddings...")
        self.embedding_model = GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-001"
        )
        
        # Text splitter (like in your notebook)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=400,
            chunk_overlap=50,
            length_function=len,
        )
        
        # Document processor
        self.doc_processor = DocumentProcessor()
        
        # ChromaDB vector store
        self.vectordb = None
        self.documents = []
        
        # Load and index knowledge base
        self._index_knowledge_base()
        
        print(f"RAGEngine initialized. Documents indexed: {len(self.documents)}")
    
    def _index_knowledge_base(self):
        """Index all documents from knowledge_base folder with persistence"""
        kb_dir = os.path.join(os.path.dirname(__file__), "knowledge_base")
        
        # Check if ChromaDB already exists (avoid re-embedding)
        if os.path.exists(CHROMA_PERSIST_DIR) and os.listdir(CHROMA_PERSIST_DIR):
            print(f"Loading existing ChromaDB from {CHROMA_PERSIST_DIR}")
            try:
                self.vectordb = Chroma(
                    persist_directory=CHROMA_PERSIST_DIR,
                    embedding_function=self.embedding_model,
                    collection_name="knowledge_base"
                )
                # Get document count
                collection = self.vectordb._collection
                doc_count = collection.count()
                self.documents = [None] * doc_count  # Placeholder for count
                print(f"Loaded {doc_count} chunks from persistent ChromaDB")
                return
            except Exception as e:
                print(f"Error loading ChromaDB: {e}, will rebuild...")
        
        if not os.path.exists(kb_dir):
            os.makedirs(kb_dir)
            print(f"Created knowledge_base directory at {kb_dir}")
            return
        
        all_documents = []
        
        # Process all supported files
        supported_extensions = ('.txt', '.md', '.pdf', '.docx')
        
        for filename in os.listdir(kb_dir):
            if filename.endswith(supported_extensions):
                filepath = os.path.join(kb_dir, filename)
                
                # Load document
                content = self.doc_processor.load_document(filepath)
                if not content:
                    continue
                
                # Create LangChain Document
                doc = Document(
                    page_content=content,
                    metadata={"source": filename}
                )
                all_documents.append(doc)
                print(f"Loaded: {filename}")
        
        # Also check for sample_kb.txt in parent directory
        sample_kb = os.path.join(os.path.dirname(__file__), "sample_kb.txt")
        if os.path.exists(sample_kb):
            content = self.doc_processor.load_document(sample_kb)
            if content:
                doc = Document(
                    page_content=content,
                    metadata={"source": "sample_kb.txt"}
                )
                all_documents.append(doc)
                print(f"Loaded: sample_kb.txt")
        
        if not all_documents:
            print("No documents found to index")
            return
        
        # Split into chunks using RecursiveCharacterTextSplitter
        print(f"Splitting {len(all_documents)} documents into chunks...")
        chunks = self.text_splitter.split_documents(all_documents)
        self.documents = chunks
        print(f"Created {len(chunks)} chunks")
        
        # Create ChromaDB vector store with Google embeddings (persistent)
        print("Creating vector embeddings with Google's embedding model...")
        print("This may take a moment due to API rate limits...")
        
        # Add delay between batches to avoid rate limits
        batch_size = 20
        for i in range(0, len(chunks), batch_size):
            if i > 0:
                time.sleep(3)  # Wait 3 seconds between batches
                print(f"  Embedding batch {i//batch_size + 1}...")
        
        self.vectordb = Chroma.from_documents(
            documents=chunks,
            embedding=self.embedding_model,
            collection_name="knowledge_base",
            persist_directory=CHROMA_PERSIST_DIR
        )
        print(f"Indexed {len(chunks)} chunks into ChromaDB (persisted)")
    
    def _retrieve_context(self, query: str, k: int = 4, threshold: float = 0.5) -> Tuple[List[Document], List[float]]:
        """
        Retrieve relevant documents using similarity search with score
        Like in your notebook: similarity_search_with_score
        
        Threshold: 0.5 means only highly relevant docs pass
        (lower score = more similar in Chroma's L2 distance)
        """
        if self.vectordb is None:
            return [], []
        
        # Perform similarity search with scores
        results = self.vectordb.similarity_search_with_score(query, k=k)
        
        # Filter by threshold - only keep truly relevant docs
        # Chroma L2 distance: 0 = identical, >1 = very different
        relevant_docs = []
        scores = []
        
        for doc, score in results:
            if score < threshold:  # Only highly relevant docs
                relevant_docs.append(doc)
                scores.append(score)
        
        return relevant_docs, scores
    
    def _calculate_confidence(self, scores: List[float], response: str, query: str, context_found: bool) -> float:
        """
        Calculate dynamic confidence score based on:
        1. Whether relevant context was found
        2. Retrieval similarity scores (lower = better in Chroma)
        3. Response quality indicators
        
        Confidence meaning:
        - 90%+: Highly confident, answer from knowledge base
        - 70-89%: Moderately confident, partial match
        - 50-69%: Low confidence, generic response
        - <50%: Very low confidence, out of scope
        """
        # No context found = low confidence general response
        if not context_found or not scores:
            return 0.35  # AI is answering without knowledge base
        
        # Convert Chroma L2 distances to similarity scores
        # L2 distance: 0 = identical, 0.5 = very similar, 1+ = different
        # More aggressive scaling for better differentiation
        similarities = []
        for score in scores:
            if score < 0.2:
                sim = 0.95  # Nearly identical
            elif score < 0.35:
                sim = 0.85  # Very similar
            elif score < 0.5:
                sim = 0.70  # Similar
            else:
                sim = max(0.3, 1 - score)  # Less similar
            similarities.append(sim)
        
        # Factor 1: Best match quality (weight: 50%)
        best_similarity = max(similarities)
        best_score = best_similarity * 0.50
        
        # Factor 2: Average relevance (weight: 30%)
        avg_similarity = sum(similarities) / len(similarities)
        avg_score = avg_similarity * 0.30
        
        # Factor 3: Response quality (weight: 20%)
        quality_score = 0.15
        
        # Penalize uncertain/awkward responses
        uncertainty_phrases = [
            "context does not contain", "i'm not sure", "i don't know",
            "i cannot", "i can't", "unclear", "not available",
            "cannot help", "outside my knowledge", "no information"
        ]
        response_lower = response.lower()
        uncertainty_count = sum(1 for phrase in uncertainty_phrases if phrase in response_lower)
        
        if uncertainty_count > 0:
            quality_score = 0.05  # Heavy penalty for uncertain responses
        elif len(response) > 150:
            quality_score = 0.20  # Bonus for detailed answers
        
        # Calculate total confidence
        total_confidence = best_score + avg_score + quality_score
        
        # Clamp between 0.15 and 0.98
        return max(0.15, min(0.98, total_confidence))
    
    def get_response(self, query: str) -> dict:
        """
        Generate AI response using RAG
        Similar to chatbot_agent in your notebook
        """
        try:
            # Retrieve relevant context
            relevant_docs, scores = self._retrieve_context(query, k=4, threshold=0.8)
            
            context_found = len(relevant_docs) > 0
            
            if context_found:
                # Build context from retrieved documents
                context = "\n\n".join([doc.page_content for doc in relevant_docs])
                sources = list(set([doc.metadata.get("source", "unknown") for doc in relevant_docs]))
                
                # Create prompt with context - natural responses only
                prompt = f"""You are a helpful customer support assistant. Answer the question using ONLY the information from the context below.

Rules:
- If the context has the answer, provide it clearly and helpfully
- If the context doesn't have the answer, give a helpful general response WITHOUT mentioning "context" or "provided information"
- Never say phrases like "the context does not contain" or "based on the provided context"
- Be conversational and professional
- If you truly can't help, suggest contacting support@company.com

Context:
{context}

Customer Question: {query}

Your helpful response:"""
                
                # Get response from LLM
                response = self.llm.invoke(prompt)
                ai_response = response.content
                
            else:
                # No relevant context - provide general helpful response
                sources = []
                prompt = f"""You are a helpful customer support assistant. The customer asked a question that may be outside your knowledge base.

Rules:
- Provide a helpful, general response if possible
- Be honest but friendly - never sound robotic
- If it's truly outside your scope, politely suggest they contact support@company.com
- Keep response concise and professional

Customer Question: {query}

Your helpful response:"""
                
                response = self.llm.invoke(prompt)
                ai_response = response.content
                scores = []  # No retrieval scores
            
            # Calculate confidence with context awareness
            confidence = self._calculate_confidence(scores, ai_response, query, context_found)
            
            return {
                "response": ai_response,
                "confidence_score": round(confidence, 2),
                "sources": sources,
                "success": True,
                "context_found": len(relevant_docs) > 0,
                "chunks_retrieved": len(relevant_docs)
            }
            
        except Exception as e:
            print(f"Error in get_response: {str(e)}")
            import traceback
            traceback.print_exc()
            
            return {
                "response": "I apologize, but I encountered an issue processing your request. Please try again or contact our support team directly.",
                "confidence_score": 0.0,
                "sources": [],
                "success": False,
                "context_found": False,
                "chunks_retrieved": 0
            }
    
    def categorize_ticket(self, description: str) -> dict:
        """Automatically categorize a support ticket using AI"""
        try:
            # Try to find similar context
            relevant_docs, scores = self._retrieve_context(description, k=3, threshold=1.0)
            
            context = ""
            if relevant_docs:
                context = f"\nSimilar past context: {relevant_docs[0].page_content[:500]}"
            
            prompt = f"""Analyze this customer support ticket and categorize it.

TICKET DESCRIPTION: {description}
{context}

Categorize into EXACTLY ONE of these categories:
- Technical Support (app issues, bugs, errors, not working)
- Billing (payments, refunds, charges, subscription costs)
- Account (login, password, profile, settings)
- General Inquiry (questions, information, how-to)
- Feature Request (suggestions, improvements, new features)

Assign priority:
- URGENT: System down, security issue, complete blocker
- HIGH: Major functionality broken, billing errors
- MEDIUM: Partial issues, inconvenience
- LOW: Questions, minor issues, feature requests

Respond in EXACTLY this format:
CATEGORY: [category name]
PRIORITY: [LOW/MEDIUM/HIGH/URGENT]
REASON: [brief reason]"""

            response = self.llm.invoke(prompt)
            response_text = response.content.strip()
            
            # Parse response
            category = "General Inquiry"
            priority = "MEDIUM"
            reason = ""
            
            for line in response_text.split('\n'):
                line = line.strip()
                if line.startswith('CATEGORY:'):
                    category = line.split(':', 1)[1].strip()
                elif line.startswith('PRIORITY:'):
                    priority = line.split(':', 1)[1].strip().upper()
                elif line.startswith('REASON:'):
                    reason = line.split(':', 1)[1].strip()
            
            # Validate
            valid_categories = ["Technical Support", "Billing", "Account", "General Inquiry", "Feature Request"]
            if category not in valid_categories:
                for vc in valid_categories:
                    if vc.lower() in category.lower():
                        category = vc
                        break
                else:
                    category = "General Inquiry"
            
            valid_priorities = ["LOW", "MEDIUM", "HIGH", "URGENT"]
            if priority not in valid_priorities:
                priority = "MEDIUM"
            
            confidence = 0.85 if scores else 0.75
            
            return {
                "category": category,
                "priority": priority,
                "confidence": round(confidence, 2),
                "reason": reason
            }
            
        except Exception as e:
            print(f"Error in categorize_ticket: {e}")
            return {
                "category": "General Inquiry",
                "priority": "MEDIUM",
                "confidence": 0.5,
                "reason": "Error during categorization"
            }
    
    def reload_knowledge_base(self, force=False):
        """Reload and re-index all documents. Set force=True to rebuild embeddings."""
        import shutil
        
        if force and os.path.exists(CHROMA_PERSIST_DIR):
            shutil.rmtree(CHROMA_PERSIST_DIR)
            print("Cleared persistent ChromaDB for full re-index")
        
        self.vectordb = None
        self.documents = []
        self._index_knowledge_base()
        return {
            "status": "success", 
            "documents_indexed": len(self.documents)
        }
    
    def get_stats(self) -> dict:
        """Get knowledge base statistics"""
        chunk_count = len(self.documents)
        if self.vectordb:
            try:
                chunk_count = self.vectordb._collection.count()
            except:
                pass
        
        return {
            "total_chunks": chunk_count,
            "embedding_model": "models/gemini-embedding-001",
            "llm_model": "gemini-2.5-flash",
            "vector_store": "ChromaDB (persistent)"
        }
